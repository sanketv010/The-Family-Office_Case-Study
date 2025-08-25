import re
import json
import docx
from pathlib import Path
from typing import Dict, Any

class FinancialDocumentParser:
    def __init__(self):
        self.target_entities = [
            "Party A", 
            "Party B", 
            "Initial Valuation Date", 
            "Notional Amount",
            "Valuation Date", 
            "Termination Date", 
            "Underlying", 
            "Coupon (C)", 
            "Barrier (B)", 
            "Business Day"
        ]

    def extract_text(self, file_path:str) -> str:
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text = text + paragraph.text + "\n"
                
            for table in doc.tables:
                for r in table.rows:
                    for c in r.cells:
                        text = text + c.text + "\n"
                        
            return text
        except Exception as e:
            print(f"Error extracting the text: {e}")
            return ""

    def clean_text(self, text:str) -> str:
        text = text.strip()
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\*+', '', text)
        return text

    def clean_extracted_value(self, value:str) -> str:
        value = re.sub(r'^[\-\*:\s]+', '', value)
        value = re.sub(r'[\-\*]+$', '', value)
        value = re.sub(r'\s+', ' ', value)
        
        if re.search(r'\*+TBD\*+', value, re.IGNORECASE):
            return "TBD"
            
        return value.strip()

    def extract_entities(self, text:str) -> Dict[str, str]:
        entities = {}
        
        # rules for extracting each entities from the text
        patterns = {
            "Party A": r"Party A[:\s]+([^\n\r]+?)(?=\s+Party B|\s+Trade Date|\n|$)",
            "Party B": r"Party B[:\s]+([^\n\r]+?)(?=\s+Trade Date|\s+Initial|\n|$)",
            "Initial Valuation Date": r"Initial Valuation\s+Date[:\s]+([^\n\r]+?)(?=\s+Effective|\s+Notional|\n|$)",
            "Notional Amount": r"Notional Amount[:\s\(N\)]*([^\n\r]+?)(?=\s+Upfront|\s+Valuation|\n|$)",
            "Valuation Date": r"(?<!Initial\s)Valuation Date[:\s]+([^\n\r]+?)(?=\s+Termination|\s+Underlying|\n|$)",
            "Termination Date": r"Termination Date[:\s]+([^\n\r]+?)(?=\s+Underlying|\s+Exchange|\n|$)",
            "Underlying": r"Underlying[:\s]+([^\n\r]+?)(?=\s+Exchange|\s+Coupon|\n|$)",
            "Coupon (C)": r"Coupon[:\s\(C\)]*([^\n\r]+?)(?=\s+Barrier|\s+Interest|\n|$)",
            "Business Day": r"Business Day[:\s]+([^\n\r]+?)(?=\s+Future|\s+Calculation|\n|$)"
        }
        
        barrier_patterns = [
            r"Barrier[:\s]*\([B]\)[:\s]*([0-9]+\.?[0-9]*%[^\n\r]*?)(?=\s+Interest|\s+##|\n|$)",
            r"Barrier[:\s]*\([B]\)[:\s]*([^\n\r]+?)(?=\s+Interest|\s+##|\n|$)",
            r"Barrier[:\s]+([0-9]+\.?[0-9]*%[^\n\r]*?)(?=\s+Interest|\s+##|\n|$)"
        ]
        
        for entity, pattern in patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                value = self.clean_extracted_value(match.group(1))
                if value:
                    entities[entity] = value
        
        for pattern in barrier_patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                value = self.clean_extracted_value(match.group(1))
                if value and len(value) < 50: 
                    entities["Barrier (B)"] = value
                    break
        
        return entities

    def parse_document(self, file_path:str) -> Dict[str, Any]:
        if not Path(file_path).exists():
            return {"error": f"File not found: {file_path}"}
        
        if file_path.lower().endswith('.docx'):
            raw_text = self.extract_text(file_path)
        else:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    raw_text = f.read()
            except Exception as e:
                return {"error": f"Error reading file: {e}"}
        
        if not raw_text:
            return {"error": "Could not extract text from document"}
        
        clean_text = self.clean_text(raw_text)
        extracted_entities = self.extract_entities(clean_text)
        
        return {
            "extracted_entities": extracted_entities,
            "entities_found": len(extracted_entities),
            "target_entities": self.target_entities
        }

    def save_results(self, results:Dict[str, Any], output_path:str = None):
        if output_path is None:
            output_path = f"ner_results.json"
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(results, f, indent=2, ensure_ascii=False)
            print(f"Results saved to: {output_path}")
        except Exception as e:
            print(f"Error saving results: {e}")

def main():
    parser = FinancialDocumentParser()
    
    doc_path = "" #provide the file path
    
    results = parser.parse_document(doc_path)
    
    if "error" in results:
        print(f"Error: {results['error']}")
    else:
        print(f"\nExtracted {results['entities_found']} entities:")
        print("-" * 50)
        
        for entity, value in results['extracted_entities'].items():
            print(f"{entity}: {value}")
        
        parser.save_results(results)
    
    print("\nProcessing complete.")

if __name__ == "__main__":
    main()



